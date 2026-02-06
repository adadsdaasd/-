"""
智能填表模块 (Smart Form Filler)
================================
功能：
1. 上传空表格/文档，AI 自动识别需要填写的字段
2. 根据用户档案智能匹配数据
3. AI 润色话术（让回答更专业）
4. 支持 Excel、Word、纯文本问题
"""

import streamlit as st
import pandas as pd
import json
import re
import io
from typing import Dict, List, Tuple, Optional
from datetime import datetime

from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ai_services import create_ai_client, DEFAULT_MODEL

from research_models import (
    load_research_profiles,
    get_research_profile_by_id,
    flatten_profile_for_template,
    get_publications_summary,
    get_grants_summary
)
from profile_validation import validate_general_profile, validate_research_profile

# ==================== 常量 ====================
# 说明：LLM base_url / model / client 统一由 ai_services.py 管理


def ai_identify_fields(client, content: str) -> List[Dict]:
    """
    AI 识别表格/文档中需要填写的字段
    
    返回: [{"field": "字段名", "type": "factual/subjective", "description": "说明"}]
    """
    prompt = """你是一个表单分析专家。请分析以下表格/文档内容，识别出所有需要填写的字段。

对于每个字段，判断它是：
1. factual（事实类）：如姓名、电话、学历等，需要准确的事实信息
2. subjective（主观类）：如自我介绍、个人优势、研究计划等，需要润色的文字描述

返回 JSON 数组格式，不要添加其他文字：
[
  {"field": "字段名", "type": "factual", "description": "简短说明"},
  {"field": "字段名", "type": "subjective", "description": "简短说明"}
]

文档内容：
"""
    
    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": content[:3000]}  # 限制长度
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        result = response.choices[0].message.content.strip()
        
        # 清理 markdown 代码块
        if result.startswith("```"):
            result = re.sub(r'^```(?:json)?\n?', '', result)
            result = re.sub(r'\n?```$', '', result)
        
        return json.loads(result)
    
    except Exception as e:
        st.error(f"AI 识别字段失败: {str(e)}")
        return []


def ai_generate_answer(client, field: str, field_type: str, profile_data: Dict, context: str = "") -> str:
    """
    AI 根据字段和用户数据生成答案（单个字段）
    """
    profile_summary = _build_profile_summary(profile_data)
    
    if field_type == "factual":
        prompt = f"""根据以下用户信息，为"{field}"字段提供准确的答案。
如果信息不存在，回复"未提供"。只返回答案，不要解释。

用户信息：
{profile_summary}

字段：{field}
答案："""
    else:
        prompt = f"""你是一个专业的文书润色专家。请根据用户信息，为"{field}"字段撰写一段专业、得体的回答。

要求：
1. 语言正式但不生硬
2. 突出优势和亮点
3. 符合学术/职场规范
4. 100-200字左右（除非是简短字段）

用户信息：
{profile_summary}

{f"补充上下文：{context}" if context else ""}

字段：{field}
专业回答："""
    
    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5 if field_type == "subjective" else 0.2,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"[生成失败: {str(e)}]"


def ai_batch_generate_answers(client, fields: List[Dict], profile_data: Dict, style: str = "professional") -> Dict[str, str]:
    """
    批量生成所有字段的答案（一次 API 调用，大幅提升速度）
    
    Args:
        client: AI 客户端
        fields: 字段列表 [{"field": "字段名", "type": "factual/subjective"}]
        profile_data: 用户档案数据
        style: 润色风格
    
    Returns:
        {"字段名": "答案", ...}
    """
    profile_summary = _build_profile_summary(profile_data)
    
    # 构建字段列表描述
    fields_desc = []
    for i, f in enumerate(fields, 1):
        field_name = f["field"]
        field_type = f.get("type", "subjective")
        type_hint = "（事实类，直接从信息中提取）" if field_type == "factual" else "（主观类，需要润色撰写）"
        fields_desc.append(f"{i}. {field_name} {type_hint}")
    
    fields_text = "\n".join(fields_desc)
    
    style_hints = {
        "professional": "专业正式，适合职场和商务场合",
        "academic": "学术规范，适合论文和学术申请",
        "friendly": "亲和友好，适合自我介绍和面试"
    }
    style_hint = style_hints.get(style, "专业正式")
    
    prompt = f"""你是一个专业的表单填写助手。请根据以下用户信息，为所有字段生成答案。

【用户信息】
{profile_summary}

【需要填写的字段】
{fields_text}

【要求】
1. 事实类字段：直接从用户信息中提取准确答案，没有则填"未提供"
2. 主观类字段：根据用户信息撰写专业、得体的回答（{style_hint}），100-200字
3. 突出用户的优势和亮点
4. 返回格式：严格按 JSON 格式返回，键为字段名，值为答案

【返回格式示例】
{{
  "姓名": "张三",
  "个人优势": "具有丰富的..."
}}

请直接返回 JSON，不要添加其他文字："""

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=3000
        )
        
        result = response.choices[0].message.content.strip()
        
        # 清理 markdown 代码块
        if result.startswith("```"):
            result = re.sub(r'^```(?:json)?\n?', '', result)
            result = re.sub(r'\n?```$', '', result)
        
        answers = json.loads(result)
        return answers
    
    except json.JSONDecodeError:
        # JSON 解析失败，返回空字典，让调用方回退到逐个生成
        return {}
    except Exception as e:
        return {"_error": f"批量生成失败: {str(e)}"}


def ai_polish_text(client, original_text: str, style: str = "professional") -> str:
    """
    AI 润色文本
    
    Args:
        original_text: 原始文本
        style: 风格 - professional/academic/friendly
    """
    style_prompts = {
        "professional": "专业正式，适合职场和商务场合",
        "academic": "学术规范，适合论文和学术申请",
        "friendly": "亲和友好，适合自我介绍和面试"
    }
    
    prompt = f"""请润色以下文本，使其{style_prompts.get(style, "更加专业得体")}。

要求：
1. 保持原意不变
2. 改善表达方式
3. 增强说服力
4. 修正语法错误

原文：
{original_text}

润色后："""
    
    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            max_tokens=1000
        )
        
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        return original_text


def _build_profile_summary(profile_data: Dict) -> str:
    """构建用户信息摘要"""
    lines = []
    
    # 基本信息
    if profile_data.get("姓名"):
        lines.append(f"姓名：{profile_data['姓名']}")
    
    contact = profile_data.get("联系方式", {})
    if isinstance(contact, dict):
        if contact.get("电话"):
            lines.append(f"电话：{contact['电话']}")
        if contact.get("邮箱"):
            lines.append(f"邮箱：{contact['邮箱']}")
    
    # 教育背景
    education = profile_data.get("education_history", [])
    if education:
        edu_str = "教育经历：" + "; ".join([
            f"{e.get('degree', '')} - {e.get('institution', '')} - {e.get('major', '')}"
            for e in education[:3]
        ])
        lines.append(edu_str)
    elif profile_data.get("教育背景"):
        lines.append(f"教育背景：{profile_data['教育背景']}")
    
    # 工作/研究经历
    experience = profile_data.get("工作经历", [])
    if experience:
        if isinstance(experience, list):
            lines.append(f"工作经历：{'; '.join(experience[:3])}")
        else:
            lines.append(f"工作经历：{experience}")
    
    # 技能
    skills = profile_data.get("技能特长", [])
    if skills:
        if isinstance(skills, list):
            lines.append(f"技能特长：{', '.join(skills)}")
        else:
            lines.append(f"技能特长：{skills}")
    
    # 论文
    publications = profile_data.get("publications", [])
    if publications:
        pub_summary = get_publications_summary(profile_data)
        lines.append(f"论文发表：共 {pub_summary['total']} 篇 (SCI: {pub_summary['sci']}, EI: {pub_summary['ei']})")
    
    # 项目
    grants = profile_data.get("grants", [])
    if grants:
        grant_summary = get_grants_summary(profile_data)
        lines.append(f"科研项目：共 {grant_summary['total']} 项，其中主持 {grant_summary['as_pi']} 项")
    
    # 个人优势
    if profile_data.get("个人优势"):
        lines.append(f"个人优势：{profile_data['个人优势']}")
    
    # 可发展方向
    dev_direction = profile_data.get("可发展方向", {})
    if isinstance(dev_direction, dict) and dev_direction.get("短期建议"):
        lines.append(f"发展方向：{dev_direction['短期建议']}")
    
    return "\n".join(lines) if lines else "暂无详细信息"


# ==================== 文件处理 ====================

def extract_excel_content(file) -> Tuple[str, pd.DataFrame]:
    """提取 Excel 内容用于 AI 分析"""
    file.seek(0)
    
    # 读取为 DataFrame
    df = pd.read_excel(file)
    
    # 转换为文本描述
    content_lines = []
    content_lines.append("表格列名：" + ", ".join(df.columns.tolist()))
    
    # 如果有数据行，展示结构
    if len(df) > 0:
        content_lines.append("\n表格结构示例：")
        for col in df.columns:
            sample = df[col].iloc[0] if pd.notna(df[col].iloc[0]) else "(空)"
            content_lines.append(f"- {col}: {sample}")
    
    return "\n".join(content_lines), df


def extract_word_content(file) -> str:
    """提取 Word 内容用于 AI 分析"""
    if not DOCX_AVAILABLE:
        return "无法读取 Word 文件（未安装 python-docx）"
    
    file.seek(0)
    doc = Document(file)
    
    content_lines = []
    
    # 提取段落
    for para in doc.paragraphs:
        if para.text.strip():
            content_lines.append(para.text)
    
    # 提取表格
    for table in doc.tables:
        content_lines.append("\n[表格内容]")
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                content_lines.append(row_text)
    
    return "\n".join(content_lines)


def parse_text_questions(text: str) -> List[str]:
    """解析纯文本问题列表"""
    lines = text.strip().split('\n')
    questions = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 移除常见序号格式
        line = re.sub(r'^[\d]+[\.、\)）]\s*', '', line)
        line = re.sub(r'^[一二三四五六七八九十]+[、\.]\s*', '', line)
        line = re.sub(r'^[-*•]\s*', '', line)
        
        if line and len(line) > 1:
            questions.append(line)
    
    return questions


# ==================== 表格模式检测 ====================

# 一表多人的特征关键词（横向排列的表头）
MULTI_PERSON_KEYWORDS = [
    "序号", "编号", "No.", "No", "#", "姓名", "名字", "电话", "手机", 
    "邮箱", "职位", "部门", "工号", "学号"
]

# 一人一表的特征关键词（纵向排列的标签）
SINGLE_PERSON_KEYWORDS = [
    "申请人", "填表人", "本人", "个人信息", "基本信息", "自我介绍",
    "个人优势", "职业规划", "研究方向", "项目简介"
]


def detect_form_mode_excel(file) -> Tuple[str, str, float]:
    """
    检测 Excel 表格的填写模式
    
    Returns:
        (mode, reason, confidence)
        mode: "batch" (一人一表) | "aggregate" (一表多人)
        reason: 判断依据说明
        confidence: 置信度 0.0-1.0
    """
    file.seek(0)
    
    try:
        df = pd.read_excel(file)
    except Exception as e:
        return "batch", f"无法解析表格: {str(e)}", 0.5
    
    columns = [str(col).strip() for col in df.columns.tolist()]
    num_columns = len(columns)
    num_rows = len(df)
    
    # 特征分数
    aggregate_score = 0  # 一表多人得分
    batch_score = 0  # 一人一表得分
    reasons = []
    
    # 检查1: 列数 - 多列通常是一表多人
    if num_columns >= 5:
        aggregate_score += 2
        reasons.append(f"列数较多({num_columns}列)")
    elif num_columns <= 2:
        batch_score += 2
        reasons.append(f"列数较少({num_columns}列)，像是标签-值结构")
    
    # 检查2: 是否有序号列
    for col in columns:
        col_lower = col.lower()
        if col_lower in ["序号", "编号", "no", "no.", "#", "id"]:
            aggregate_score += 3
            reasons.append(f"存在序号列 '{col}'")
            break
    
    # 检查3: 列名是否包含多人特征关键词
    multi_keyword_count = 0
    for col in columns:
        for keyword in MULTI_PERSON_KEYWORDS:
            if keyword in col:
                multi_keyword_count += 1
                break
    
    if multi_keyword_count >= 3:
        aggregate_score += 3
        reasons.append(f"列名包含多个人员信息字段({multi_keyword_count}个)")
    
    # 检查4: 是否有空数据行（等待填写）
    empty_rows = df.isna().all(axis=1).sum()
    if empty_rows >= 2:
        aggregate_score += 2
        reasons.append(f"存在{empty_rows}个空行待填写")
    
    # 检查5: 第一列是否像是标签（一人一表特征）
    if num_columns == 2:
        first_col_values = df.iloc[:, 0].dropna().astype(str).tolist()
        label_like_count = 0
        for val in first_col_values[:10]:
            # 检查是否像标签（较短、包含冒号或问号）
            if len(val) < 20 and (":" in val or "：" in val or "?" in val or "？" in val):
                label_like_count += 1
            # 检查是否包含一人一表关键词
            for keyword in SINGLE_PERSON_KEYWORDS:
                if keyword in val:
                    label_like_count += 1
                    break
        
        if label_like_count >= 3:
            batch_score += 3
            reasons.append("第一列像是表单标签")
    
    # 检查6: 行数判断
    if num_rows >= 5 and num_columns >= 3:
        aggregate_score += 1
        reasons.append(f"表格有{num_rows}行，适合填入多人")
    elif num_rows >= 10 and num_columns == 2:
        batch_score += 1
        reasons.append(f"纵向结构，{num_rows}个字段")
    
    # 计算最终结果
    total_score = aggregate_score + batch_score
    if total_score == 0:
        return "batch", "无法确定，默认使用一人一表", 0.5
    
    if aggregate_score > batch_score:
        confidence = min(0.95, 0.5 + (aggregate_score - batch_score) * 0.1)
        mode = "aggregate"
        mode_desc = "一表多人"
    else:
        confidence = min(0.95, 0.5 + (batch_score - aggregate_score) * 0.1)
        mode = "batch"
        mode_desc = "一人一表"
    
    reason_text = f"判断为{mode_desc}：" + "；".join(reasons[:3])
    return mode, reason_text, confidence


def detect_form_mode_word(file) -> Tuple[str, str, float]:
    """
    检测 Word 文档的填写模式
    
    Returns:
        (mode, reason, confidence)
    """
    if not DOCX_AVAILABLE:
        return "batch", "无法解析 Word 文件", 0.5
    
    file.seek(0)
    
    try:
        doc = Document(file)
    except Exception as e:
        return "batch", f"无法解析文档: {str(e)}", 0.5
    
    aggregate_score = 0
    batch_score = 0
    reasons = []
    
    # 检查表格
    tables = doc.tables
    if tables:
        for table in tables:
            num_rows = len(table.rows)
            num_cols = len(table.columns) if table.rows else 0
            
            # 多行多列的表格 -> 一表多人
            if num_rows >= 3 and num_cols >= 3:
                aggregate_score += 3
                reasons.append(f"包含 {num_rows}x{num_cols} 的表格")
            
            # 两列表格 -> 可能是一人一表
            elif num_cols == 2 and num_rows >= 5:
                batch_score += 2
                reasons.append("包含两列表格（标签-值结构）")
            
            # 检查表头
            if table.rows:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                header_text = " ".join(header_cells)
                
                # 检查是否有多人特征关键词
                multi_count = sum(1 for kw in MULTI_PERSON_KEYWORDS if kw in header_text)
                if multi_count >= 2:
                    aggregate_score += 2
                    reasons.append("表头包含多人信息字段")
    else:
        # 没有表格，检查段落
        batch_score += 1
        reasons.append("无表格，可能是文本型表单")
    
    # 检查段落中的占位符
    placeholder_count = 0
    for para in doc.paragraphs:
        text = para.text
        # 检查 {{xxx}} 占位符
        placeholders = re.findall(r'\{\{[^}]+\}\}', text)
        placeholder_count += len(placeholders)
    
    if placeholder_count > 0:
        batch_score += 2
        reasons.append(f"包含 {placeholder_count} 个占位符")
    
    # 计算结果
    total_score = aggregate_score + batch_score
    if total_score == 0:
        return "batch", "无法确定，默认使用一人一表", 0.5
    
    if aggregate_score > batch_score:
        confidence = min(0.95, 0.5 + (aggregate_score - batch_score) * 0.1)
        mode = "aggregate"
        mode_desc = "一表多人"
    else:
        confidence = min(0.95, 0.5 + (batch_score - aggregate_score) * 0.1)
        mode = "batch"
        mode_desc = "一人一表"
    
    reason_text = f"判断为{mode_desc}：" + "；".join(reasons[:3]) if reasons else f"判断为{mode_desc}"
    return mode, reason_text, confidence


def detect_form_mode(file, file_type: str) -> Tuple[str, str, float]:
    """
    检测表格填写模式（统一入口）
    
    Args:
        file: 上传的文件对象
        file_type: "excel" | "word"
    
    Returns:
        (mode, reason, confidence)
        mode: "batch" (一人一表) | "aggregate" (一表多人)
        reason: 判断依据说明
        confidence: 置信度 0.0-1.0
    """
    if file_type == "excel":
        return detect_form_mode_excel(file)
    elif file_type == "word":
        return detect_form_mode_word(file)
    else:
        return "batch", "未知文件类型，默认使用一人一表", 0.5


# ==================== 填写和导出 ====================

def fill_excel_with_answers(template_file, answers: Dict[str, str]) -> bytes:
    """将答案填入 Excel 并返回"""
    template_file.seek(0)
    wb = load_workbook(template_file)
    ws = wb.active
    
    # 查找并填写
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                cell_text = cell.value.strip()
                
                # 检查是否是字段名
                for field, answer in answers.items():
                    if field in cell_text or cell_text in field:
                        # 找到下一个单元格或同行右侧单元格填写答案
                        next_cell = ws.cell(row=cell.row, column=cell.column + 1)
                        if not next_cell.value:
                            next_cell.value = answer
                            next_cell.alignment = Alignment(wrap_text=True)
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def fill_word_with_answers(template_file, answers: Dict[str, str]) -> bytes:
    """将答案填入 Word 并返回"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装")
    
    template_file.seek(0)
    doc = Document(template_file)
    
    # 填写段落中的字段
    for para in doc.paragraphs:
        for field, answer in answers.items():
            if field in para.text:
                # 在字段后添加答案
                para.text = para.text.replace(field, f"{field}：{answer}")
    
    # 填写表格
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                if cell_text in answers:
                    # 如果有下一个单元格，填入答案
                    if i + 1 < len(cells):
                        cells[i + 1].text = answers[cell_text]
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def export_answers_to_csv(answers: Dict[str, str]) -> str:
    """导出答案为 CSV"""
    df = pd.DataFrame([
        {"字段": k, "答案": v} for k, v in answers.items()
    ])
    return df.to_csv(index=False, encoding='utf-8-sig')


# ==================== Streamlit UI ====================

def render_smart_form_filler(api_key: str):
    """渲染智能填表界面"""
    
    st.header("🪄 智能填表")
    st.markdown("上传空表格，AI 自动识别字段并填写，支持润色话术")
    
    if not api_key:
        st.warning("⚠️ 请先在侧边栏输入 DeepSeek API Key")
        return
    
    st.markdown("---")
    
    # 选择输入方式
    input_method = st.radio(
        "选择输入方式",
        ["📄 上传 Excel 表格", "📝 上传 Word 文档", "✍️ 粘贴问题列表"],
        horizontal=True
    )
    
    # 选择用户档案（根据当前用户模式过滤）
    st.markdown("---")
    st.subheader("👤 选择要使用的档案")
    
    from self_config import load_self_profile_from_orgstore, get_self_person_id
    from store_org import load_profiles_multi, load_groups, load_people
    
    # 获取当前用户模式
    current_mode = st.session_state.get('mode', 'single')
    
    all_profiles = []
    
    if current_mode == 'single':
        # 个人版：从 OrgStore 加载「我自己」的档案（新方式：个人版=多人版中的一员）
        self_data = load_self_profile_from_orgstore()
        if self_data:
            profile_data = self_data.get("profile", {})
            person_id = self_data.get("person_id", "single")
            if isinstance(profile_data, dict):
                all_profiles.append({
                    "id": person_id,
                    "name": self_data.get("name") or profile_data.get("姓名", "个人版用户"),
                    "source": "个人档案",
                    "data": profile_data
                })
        else:
            # 未绑定「我自己」，提示用户
            st.info("💡 请先在「数据管理」中设置个人信息或绑定「我是谁」")
    else:
        # 多人版：加载多人版档案
        multi_profiles = load_profiles_multi()
        for mp in multi_profiles:
            profile_data = mp.get("profile", {})
            if isinstance(profile_data, dict):
                all_profiles.append({
                    "id": mp["id"],
                    "name": profile_data.get("姓名", mp.get("name", "未知")),
                    "source": "团队成员",
                    "data": profile_data,
                    "groups": mp.get("groups", []),  # 保存所属团队信息
                    "memberships": mp.get("memberships", [])
                })
    
    # 同时加载研究档案（两种模式都可用）
    research_profiles = load_research_profiles()
    for p in research_profiles:
        all_profiles.append({
            "id": p["id"],
            "name": p.get("姓名", "未知"),
            "source": "科研档案",
            "data": p
        })
    
    if not all_profiles:
        mode_name = "个人版" if current_mode == 'single' else "多人版"
        st.warning(f"没有可用的档案，请先在「数据管理」中添加{mode_name}用户信息")
        return
    
    profile_options = {p["id"]: f"{p['name']} ({p['source']})" for p in all_profiles}
    
    # 多人版支持多选：按人员选择 或 按团队选择
    if current_mode == 'multi' and len(all_profiles) > 1:
        # 获取所有团队
        groups = load_groups()
        
        # 选择方式
        selection_mode = st.radio(
            "选择方式",
            ["👤 按人员选择", "👥 按团队选择"],
            horizontal=True,
            key="profile_selection_mode"
        )
        
        if "按团队选择" in selection_mode and groups:
            # 按团队选择
            st.info("💡 选择团队后，将自动选中该团队的所有成员")
            
            group_options = {g["id"]: g["name"] for g in groups}
            selected_group_ids = st.multiselect(
                "选择团队（可多选）",
                options=list(group_options.keys()),
                format_func=lambda x: f"📁 {group_options[x]}",
                key="select_groups_for_form"
            )
            
            # 根据选中的团队筛选成员
            if selected_group_ids:
                selected_profiles = []
                for p in all_profiles:
                    # 检查该成员是否属于选中的任一团队
                    memberships = p.get("memberships", [])
                    for ms in memberships:
                        if ms.get("group_id") in selected_group_ids:
                            if p not in selected_profiles:
                                selected_profiles.append(p)
                            break
                
                # 显示已选团队的成员数
                if selected_profiles:
                    group_names = [group_options.get(gid, gid) for gid in selected_group_ids]
                    st.success(f"已选择团队：{', '.join(group_names)}，共 {len(selected_profiles)} 人")
                else:
                    st.warning("选中的团队暂无成员")
                    selected_profiles = []
            else:
                selected_profiles = []
                st.info("请选择至少一个团队")
        else:
            # 按人员选择
            if "按团队选择" in selection_mode and not groups:
                st.info("💡 暂无团队数据，请先在「数据管理」中创建团队")
            
            st.info("💡 多人版模式：可选择多个人员批量生成")
            selected_profile_ids = st.multiselect(
                "选择档案（可多选）",
                options=list(profile_options.keys()),
                default=[list(profile_options.keys())[0]] if profile_options else [],
                format_func=lambda x: profile_options[x],
                key="select_profiles_for_form"
            )
            selected_profiles = [p for p in all_profiles if p["id"] in selected_profile_ids]
    else:
        selected_profile_id = st.selectbox(
            "选择档案",
            options=list(profile_options.keys()),
            format_func=lambda x: profile_options[x]
        )
        selected_profiles = [p for p in all_profiles if p["id"] == selected_profile_id]
    
    if not selected_profiles:
        st.warning("请至少选择一个档案")
        return
    
    # 显示已选档案
    if len(selected_profiles) == 1:
        with st.expander("查看档案信息"):
            st.json(selected_profiles[0]["data"])
    else:
        with st.expander(f"查看已选 {len(selected_profiles)} 人的档案"):
            for p in selected_profiles:
                st.markdown(f"**{p['name']}** ({p['source']})")
                st.json(p["data"])
                st.markdown("---")
    
    st.markdown("---")
    
    # 根据输入方式处理
    if "Excel" in input_method:
        _handle_excel_input(api_key, selected_profiles, all_profiles)
    elif "Word" in input_method:
        _handle_word_input(api_key, selected_profiles, all_profiles)
    else:
        _handle_text_input(api_key, selected_profiles)


def _handle_excel_input(api_key: str, selected_profiles: List[Dict], all_profiles: List[Dict]):
    """处理 Excel 输入（支持模式检测和多人批量）"""
    
    uploaded_file = st.file_uploader(
        "上传 Excel 表格",
        type=['xlsx', 'xls'],
        help="上传需要填写的空白表格"
    )
    
    if not uploaded_file:
        st.info("请上传 Excel 文件")
        return
    
    # 提取内容
    content, df = extract_excel_content(uploaded_file)
    
    st.markdown("**表格预览：**")
    st.dataframe(df, use_container_width=True)
    
    # 自动检测表格模式
    st.markdown("---")
    st.subheader("📊 表格模式检测")
    
    detected_mode, reason, confidence = detect_form_mode(uploaded_file, "excel")
    
    # 显示检测结果
    mode_labels = {"batch": "一人一表", "aggregate": "一表多人"}
    confidence_pct = int(confidence * 100)
    
    st.markdown(f"**AI 判断结果：** {mode_labels[detected_mode]} (置信度: {confidence_pct}%)")
    st.caption(reason)
    
    # 手动切换选项
    form_mode = st.radio(
        "选择填写模式",
        ["batch", "aggregate"],
        index=0 if detected_mode == "batch" else 1,
        format_func=lambda x: "📂 一人一表 (每人生成一个文件)" if x == "batch" else "📑 一表多人 (所有人填入同一表格)",
        horizontal=True,
        help="如果 AI 判断不准确，可以手动切换"
    )
    
    if form_mode != detected_mode:
        st.info("已切换为手动选择模式")
    
    # 根据模式显示不同提示
    if form_mode == "batch":
        st.success(f"📂 将为 {len(selected_profiles)} 人各生成一份填写结果")
    else:
        st.success(f"📑 将把 {len(selected_profiles)} 人的信息填入同一表格")
    
    st.markdown("---")
    
    # AI 识别字段
    if st.button("🔍 AI 识别字段", type="primary"):
        with st.spinner("AI 正在分析表格结构..."):
            client = create_ai_client(api_key)
            fields = ai_identify_fields(client, content)
        
        if fields:
            st.session_state['identified_fields'] = fields
            st.session_state['uploaded_file'] = uploaded_file
            st.session_state['form_mode'] = form_mode
            st.success(f"✅ 识别到 {len(fields)} 个字段")
    
    # 显示识别结果并生成答案
    if 'identified_fields' in st.session_state:
        current_form_mode = st.session_state.get('form_mode', form_mode)
        _render_field_filling_multi(api_key, selected_profiles, st.session_state['identified_fields'], "excel", current_form_mode)


def _handle_word_input(api_key: str, selected_profiles: List[Dict], all_profiles: List[Dict]):
    """处理 Word 输入（支持模式检测和多人批量）"""
    
    if not DOCX_AVAILABLE:
        st.error("未安装 python-docx，无法处理 Word 文件")
        return
    
    uploaded_file = st.file_uploader(
        "上传 Word 文档",
        type=['docx'],
        help="上传需要填写的文档模板"
    )
    
    if not uploaded_file:
        st.info("请上传 Word 文件")
        return
    
    # 提取内容
    content = extract_word_content(uploaded_file)
    
    with st.expander("文档内容预览"):
        st.text(content[:2000] + ("..." if len(content) > 2000 else ""))
    
    # 自动检测文档模式
    st.markdown("---")
    st.subheader("📊 文档模式检测")
    
    detected_mode, reason, confidence = detect_form_mode(uploaded_file, "word")
    
    # 显示检测结果
    mode_labels = {"batch": "一人一表", "aggregate": "一表多人"}
    confidence_pct = int(confidence * 100)
    
    st.markdown(f"**AI 判断结果：** {mode_labels[detected_mode]} (置信度: {confidence_pct}%)")
    st.caption(reason)
    
    # 手动切换选项
    form_mode = st.radio(
        "选择填写模式",
        ["batch", "aggregate"],
        index=0 if detected_mode == "batch" else 1,
        format_func=lambda x: "📂 一人一表 (每人生成一个文件)" if x == "batch" else "📑 一表多人 (所有人填入同一文档)",
        horizontal=True,
        help="如果 AI 判断不准确，可以手动切换",
        key="word_form_mode"
    )
    
    if form_mode != detected_mode:
        st.info("已切换为手动选择模式")
    
    # 根据模式显示不同提示
    if form_mode == "batch":
        st.success(f"📂 将为 {len(selected_profiles)} 人各生成一份填写结果")
    else:
        st.success(f"📑 将把 {len(selected_profiles)} 人的信息填入同一文档")
    
    st.markdown("---")
    
    # AI 识别字段
    if st.button("🔍 AI 识别字段", type="primary", key="word_identify"):
        with st.spinner("AI 正在分析文档结构..."):
            client = create_ai_client(api_key)
            fields = ai_identify_fields(client, content)
        
        if fields:
            st.session_state['identified_fields'] = fields
            st.session_state['uploaded_file'] = uploaded_file
            st.session_state['form_mode'] = form_mode
            st.success(f"✅ 识别到 {len(fields)} 个字段")
    
    if 'identified_fields' in st.session_state:
        current_form_mode = st.session_state.get('form_mode', form_mode)
        _render_field_filling_multi(api_key, selected_profiles, st.session_state['identified_fields'], "word", current_form_mode)


def _handle_text_input(api_key: str, profiles: List[Dict]):
    """处理纯文本问题输入（支持单人/多人）"""
    if not profiles:
        st.warning("请先选择一个档案")
        return
    
    questions_text = st.text_area(
        "粘贴问题列表",
        height=200,
        key="text_questions_input",
        placeholder="示例：\n1. 请介绍一下你自己\n2. 你的研究方向是什么\n3. 为什么选择我们学校"
    )
    
    if not questions_text.strip():
        st.info("请输入问题列表")
        return
    
    questions = parse_text_questions(questions_text)
    if not questions:
        st.warning("未识别到有效问题，请检查输入格式（支持换行/序号/项目符号）")
        return
    
    st.markdown(f"**识别到 {len(questions)} 个问题：**")
    for i, q in enumerate(questions, 1):
        st.markdown(f"{i}. {q}")
    
    # 转换为字段格式
    fields = [{"field": q, "type": "subjective", "description": "问答题"} for q in questions]

    st.markdown("---")
    st.subheader("📝 生成与编辑")

    # 多人：支持“每人一份 / 汇总对照表”
    if len(profiles) > 1:
        form_mode = st.radio(
            "输出方式",
            ["batch", "aggregate"],
            index=0,
            format_func=lambda x: "👤 每人一份（推荐）" if x == "batch" else "📊 汇总对照表（便于横向比较）",
            horizontal=True,
            key="text_output_mode",
        )
        _render_field_filling_multi(api_key, profiles, fields, "text", form_mode)
    else:
        _render_field_filling(api_key, profiles[0], fields, "text")


def _render_field_filling(api_key: str, profile: Dict, fields: List[Dict], file_type: str):
    """渲染字段填写界面"""
    
    st.markdown("---")
    st.subheader("📝 字段填写")

    # ========== 信息完整性校验：不完整则禁止生成/导出 ==========
    validation = None
    try:
        source_label = str(profile.get("source", ""))
        if "科研" in source_label:
            validation = validate_research_profile(profile.get("data"))
        else:
            validation = validate_general_profile(profile.get("data"))
    except Exception:
        validation = {"is_complete": False, "missing_required": ["画像结构异常"], "issues": ["校验器异常"]}

    can_output = bool(validation and validation.get("is_complete"))

    if not can_output:
        missing = validation.get("missing_required", []) if isinstance(validation, dict) else []
        st.error("⚠️ 个人必填信息不完整：无法生成/导出表格结果。"
                 + (f" 缺少：{'、'.join(missing)}" if missing else ""))
        with st.expander("查看必填项详情"):
            items_required = validation.get("items_required", []) if isinstance(validation, dict) else []
            for it in items_required:
                ok = it.get("ok", False)
                label = it.get("label", "")
                st.markdown(("✅ " if ok else "❌ ") + f"**{label}**")
    
    # 初始化答案存储
    if 'generated_answers' not in st.session_state:
        st.session_state['generated_answers'] = {}
    
    # 显示字段列表
    answers = st.session_state['generated_answers']
    
    # 润色风格选择
    polish_style = st.selectbox(
        "润色风格",
        ["professional", "academic", "friendly"],
        format_func=lambda x: {"professional": "专业正式", "academic": "学术规范", "friendly": "亲和友好"}[x]
    )
    
    # 批量生成按钮
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🚀 批量生成所有答案", use_container_width=True, disabled=not can_output):
            client = create_ai_client(api_key)
            
            with st.spinner("AI 正在批量生成所有答案...（一次调用，速度更快）"):
                # 使用批量生成（一次 API 调用）
                batch_answers = ai_batch_generate_answers(client, fields, profile["data"], polish_style)
            
            # 检查是否成功
            if "_error" in batch_answers:
                st.error(batch_answers["_error"])
            elif batch_answers:
                # 批量生成成功
                for field_info in fields:
                    field = field_info["field"]
                    if field in batch_answers:
                        answers[field] = batch_answers[field]
                
                st.session_state['generated_answers'] = answers
                st.success(f"✅ 已批量生成 {len(batch_answers)} 个字段的答案！")
                st.rerun()
            else:
                # 批量生成失败，回退到逐个生成
                st.warning("批量生成失败，正在逐个生成...")
                progress = st.progress(0)
                
                for i, field_info in enumerate(fields):
                    field = field_info["field"]
                    field_type = field_info.get("type", "subjective")
                    
                    with st.spinner(f"生成: {field}..."):
                        answer = ai_generate_answer(client, field, field_type, profile["data"])
                        answers[field] = answer
                    
                    progress.progress((i + 1) / len(fields))
                
                st.session_state['generated_answers'] = answers
                st.success("✅ 所有答案已生成！")
                st.rerun()
    
    with col2:
        if st.button("🗑️ 清空所有答案", use_container_width=True):
            st.session_state['generated_answers'] = {}
            st.rerun()
    
    st.markdown("---")
    
    # 显示每个字段的答案（可编辑）
    for field_info in fields:
        field = field_info["field"]
        field_type = field_info.get("type", "subjective")
        desc = field_info.get("description", "")
        
        type_badge = "🔵 事实类" if field_type == "factual" else "🟢 主观类"
        
        st.markdown(f"**{field}** {type_badge}")
        if desc:
            st.caption(desc)
        
        # 可编辑的答案
        current_answer = answers.get(field, "")
        new_answer = st.text_area(
            f"答案",
            value=current_answer,
            key=f"answer_{field}",
            height=100 if field_type == "subjective" else 50,
            label_visibility="collapsed"
        )
        
        # 更新答案
        if new_answer != current_answer:
            answers[field] = new_answer
            st.session_state['generated_answers'] = answers
        
        # 单独生成/润色按钮
        col_a, col_b = st.columns([1, 1])
        with col_a:
            if st.button(f"生成", key=f"gen_{field}", disabled=not can_output):
                client = create_ai_client(api_key)
                with st.spinner("生成中..."):
                    answer = ai_generate_answer(client, field, field_type, profile["data"])
                    if field_type == "subjective":
                        answer = ai_polish_text(client, answer, polish_style)
                    answers[field] = answer
                    st.session_state['generated_answers'] = answers
                st.rerun()
        
        with col_b:
            if new_answer and st.button(f"润色", key=f"polish_{field}", disabled=not can_output):
                client = create_ai_client(api_key)
                with st.spinner("润色中..."):
                    polished = ai_polish_text(client, new_answer, polish_style)
                    answers[field] = polished
                    st.session_state['generated_answers'] = answers
                st.rerun()
        
        st.markdown("---")
    
    # 导出按钮
    if answers:
        st.subheader("📤 导出结果")

        if not can_output:
            st.warning("信息不完整：导出已禁用。请先补全必填信息。")
            return
        
        col_exp1, col_exp2 = st.columns(2)
        
        with col_exp1:
            # 导出为 CSV
            csv_data = export_answers_to_csv(answers)
            st.download_button(
                label="📥 下载 CSV",
                data=csv_data,
                file_name="filled_form.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        with col_exp2:
            # 导出为 JSON
            json_data = json.dumps(answers, ensure_ascii=False, indent=2)
            st.download_button(
                label="📥 下载 JSON",
                data=json_data,
                file_name="filled_form.json",
                mime="application/json",
                use_container_width=True
            )


def _render_field_filling_multi(api_key: str, profiles: List[Dict], fields: List[Dict], file_type: str, form_mode: str):
    """渲染多人字段填写界面（支持一人一表和一表多人）"""
    
    st.markdown("---")
    
    if form_mode == "batch":
        st.subheader(f"📝 批量填写 ({len(profiles)} 人)")
        _render_batch_mode(api_key, profiles, fields, file_type)
    else:
        st.subheader(f"📝 聚合填写 ({len(profiles)} 人)")
        _render_aggregate_mode(api_key, profiles, fields, file_type)


def _render_batch_mode(api_key: str, profiles: List[Dict], fields: List[Dict], file_type: str):
    """渲染一人一表模式（每人生成独立结果）"""
    
    # 初始化多人答案存储
    if 'multi_answers' not in st.session_state:
        st.session_state['multi_answers'] = {}
    
    multi_answers = st.session_state['multi_answers']
    
    # 润色风格选择
    polish_style = st.selectbox(
        "润色风格",
        ["professional", "academic", "friendly"],
        format_func=lambda x: {"professional": "专业正式", "academic": "学术规范", "friendly": "亲和友好"}[x],
        key="batch_polish_style"
    )

    # ========== 信息完整性校验：任一人不完整则禁止生成/导出 ==========
    incomplete = []
    for p in profiles:
        try:
            src = str(p.get("source", ""))
            res = validate_research_profile(p.get("data")) if "科研" in src else validate_general_profile(p.get("data"))
        except Exception:
            res = {"is_complete": False, "missing_required": ["画像结构异常"]}
        if not res.get("is_complete"):
            incomplete.append({"name": p.get("name", "未知"), "missing": res.get("missing_required", [])})

    can_output = len(incomplete) == 0
    if not can_output:
        names = [x["name"] for x in incomplete]
        st.error("⚠️ 以下人员必填信息不完整：无法批量生成/导出。"
                 + f" 人员：{', '.join(names)}")
        with st.expander("查看缺失项"):
            for x in incomplete:
                missing = "、".join(x.get("missing", []) or [])
                st.markdown(f"- **{x['name']}**：{missing if missing else '缺失必填项'}")
    
    # 批量生成所有人的答案（优化：每人一次 API 调用）
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🚀 批量生成所有人的答案", use_container_width=True, key="batch_generate_all", disabled=not can_output):
            client = create_ai_client(api_key)
            progress = st.progress(0)
            
            for i, profile in enumerate(profiles):
                profile_id = profile["id"]
                profile_name = profile["name"]
                
                with st.spinner(f"正在为 {profile_name} 批量生成答案..."):
                    # 每人一次 API 调用
                    batch_answers = ai_batch_generate_answers(client, fields, profile["data"], polish_style)
                    
                    if batch_answers and "_error" not in batch_answers:
                        multi_answers[profile_id] = batch_answers
                    else:
                        # 回退到逐个生成
                        multi_answers[profile_id] = {}
                        for field_info in fields:
                            field = field_info["field"]
                            field_type = field_info.get("type", "subjective")
                            answer = ai_generate_answer(client, field, field_type, profile["data"])
                            multi_answers[profile_id][field] = answer
                
                progress.progress((i + 1) / len(profiles))
            
            st.session_state['multi_answers'] = multi_answers
            st.success(f"✅ 已为 {len(profiles)} 人生成答案！")
            st.rerun()
    
    with col2:
        if st.button("🗑️ 清空所有答案", use_container_width=True, key="batch_clear_all"):
            st.session_state['multi_answers'] = {}
            st.rerun()
    
    st.markdown("---")
    
    # 为每个人显示结果（使用 tabs）
    if profiles:
        person_tabs = st.tabs([p["name"] for p in profiles])
        
        for i, (tab, profile) in enumerate(zip(person_tabs, profiles)):
            with tab:
                profile_id = profile["id"]
                person_answers = multi_answers.get(profile_id, {})
                
                st.markdown(f"**{profile['name']}** ({profile['source']})")
                
                # 显示每个字段
                for field_info in fields:
                    field = field_info["field"]
                    field_type = field_info.get("type", "subjective")
                    
                    type_badge = "🔵" if field_type == "factual" else "🟢"
                    st.markdown(f"{type_badge} **{field}**")
                    
                    current_answer = person_answers.get(field, "")
                    new_answer = st.text_area(
                        f"答案",
                        value=current_answer,
                        key=f"batch_{profile_id}_{field}",
                        height=80,
                        label_visibility="collapsed"
                    )
                    
                    if new_answer != current_answer:
                        if profile_id not in multi_answers:
                            multi_answers[profile_id] = {}
                        multi_answers[profile_id][field] = new_answer
                        st.session_state['multi_answers'] = multi_answers
                
                # 单人导出
                if person_answers:
                    st.markdown("---")
                    if can_output:
                        csv_data = export_answers_to_csv(person_answers)
                        st.download_button(
                            label=f"📥 下载 {profile['name']} 的结果 (CSV)",
                            data=csv_data,
                            file_name=f"filled_form_{profile['name']}.csv",
                            mime="text/csv",
                            key=f"download_{profile_id}"
                        )
                    else:
                        st.warning("信息不完整：导出已禁用")
    
    # 批量导出所有人
    if multi_answers:
        st.markdown("---")
        st.subheader("📤 批量导出")
        
        # 合并所有人的结果为一个 JSON
        all_results = {}
        for profile in profiles:
            profile_id = profile["id"]
            if profile_id in multi_answers:
                all_results[profile["name"]] = multi_answers[profile_id]
        
        if all_results:
            if can_output:
                json_data = json.dumps(all_results, ensure_ascii=False, indent=2)
                st.download_button(
                    label="📥 下载所有人的结果 (JSON)",
                    data=json_data,
                    file_name="all_filled_forms.json",
                    mime="application/json",
                    use_container_width=True
                )
            else:
                st.warning("信息不完整：导出已禁用")


def _render_aggregate_mode(api_key: str, profiles: List[Dict], fields: List[Dict], file_type: str):
    """渲染一表多人模式（所有人填入同一表格）"""
    
    # 初始化聚合答案存储
    if 'aggregate_answers' not in st.session_state:
        st.session_state['aggregate_answers'] = {}
    
    aggregate_answers = st.session_state['aggregate_answers']
    
    # 润色风格选择
    polish_style = st.selectbox(
        "润色风格",
        ["professional", "academic", "friendly"],
        format_func=lambda x: {"professional": "专业正式", "academic": "学术规范", "friendly": "亲和友好"}[x],
        key="aggregate_polish_style"
    )

    # ========== 信息完整性校验：任一人不完整则禁止生成/导出 ==========
    incomplete = []
    for p in profiles:
        try:
            src = str(p.get("source", ""))
            res = validate_research_profile(p.get("data")) if "科研" in src else validate_general_profile(p.get("data"))
        except Exception:
            res = {"is_complete": False, "missing_required": ["画像结构异常"]}
        if not res.get("is_complete"):
            incomplete.append({"name": p.get("name", "未知"), "missing": res.get("missing_required", [])})

    can_output = len(incomplete) == 0
    if not can_output:
        names = [x["name"] for x in incomplete]
        st.error("⚠️ 以下人员必填信息不完整：无法批量生成/导出。"
                 + f" 人员：{', '.join(names)}")
        with st.expander("查看缺失项"):
            for x in incomplete:
                missing = "、".join(x.get("missing", []) or [])
                st.markdown(f"- **{x['name']}**：{missing if missing else '缺失必填项'}")
    
    # 批量生成
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🚀 批量生成所有人的答案", use_container_width=True, key="aggregate_generate_all", disabled=not can_output):
            client = create_ai_client(api_key)
            progress = st.progress(0)
            
            for i, profile in enumerate(profiles):
                profile_id = profile["id"]
                
                with st.spinner(f"正在为 {profile['name']} 批量生成答案..."):
                    # 每人一次 API 调用
                    batch_answers = ai_batch_generate_answers(client, fields, profile["data"], polish_style)
                    
                    if batch_answers and "_error" not in batch_answers:
                        aggregate_answers[profile_id] = {"_name": profile["name"], **batch_answers}
                    else:
                        # 回退到逐个生成
                        aggregate_answers[profile_id] = {"_name": profile["name"]}
                        for field_info in fields:
                            field = field_info["field"]
                            field_type = field_info.get("type", "subjective")
                            answer = ai_generate_answer(client, field, field_type, profile["data"])
                            aggregate_answers[profile_id][field] = answer
                
                progress.progress((i + 1) / len(profiles))
            
            st.session_state['aggregate_answers'] = aggregate_answers
            st.success(f"✅ 已为 {len(profiles)} 人生成答案！")
            st.rerun()
    
    with col2:
        if st.button("🗑️ 清空所有答案", use_container_width=True, key="aggregate_clear_all"):
            st.session_state['aggregate_answers'] = {}
            st.rerun()
    
    st.markdown("---")
    
    # 以表格形式展示所有人的答案
    if aggregate_answers:
        st.subheader("📊 汇总表格")
        
        # 构建 DataFrame
        table_data = []
        for profile in profiles:
            profile_id = profile["id"]
            if profile_id in aggregate_answers:
                row = {"姓名": profile["name"]}
                for field_info in fields:
                    field = field_info["field"]
                    row[field] = aggregate_answers[profile_id].get(field, "")
                table_data.append(row)
        
        if table_data:
            df = pd.DataFrame(table_data)
            
            # 使用 data_editor 允许编辑
            edited_df = st.data_editor(
                df,
                use_container_width=True,
                num_rows="fixed",
                key="aggregate_table_editor"
            )
            
            # 更新编辑后的数据
            for i, profile in enumerate(profiles):
                if i < len(edited_df):
                    profile_id = profile["id"]
                    if profile_id not in aggregate_answers:
                        aggregate_answers[profile_id] = {"_name": profile["name"]}
                    for field_info in fields:
                        field = field_info["field"]
                        if field in edited_df.columns:
                            aggregate_answers[profile_id][field] = edited_df.iloc[i][field]
            
            st.session_state['aggregate_answers'] = aggregate_answers
            
            # 导出
            st.markdown("---")
            st.subheader("📤 导出汇总表")

            if not can_output:
                st.warning("信息不完整：导出已禁用。请先补全必填信息。")
                return
            
            col_exp1, col_exp2 = st.columns(2)
            with col_exp1:
                csv_data = df.to_csv(index=False, encoding='utf-8-sig')
                st.download_button(
                    label="📥 下载汇总表 (CSV)",
                    data=csv_data,
                    file_name="aggregated_form.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col_exp2:
                # 转为 Excel
                output = io.BytesIO()
                df.to_excel(output, index=False, engine='openpyxl')
                output.seek(0)
                st.download_button(
                    label="📥 下载汇总表 (Excel)",
                    data=output.getvalue(),
                    file_name="aggregated_form.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
    else:
        # 显示空表格预览
        st.info("点击「批量生成」后，将在此显示所有人的汇总表格")
