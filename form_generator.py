"""
表单生成引擎 (Form Generator)
==============================
支持 Excel 和 Word 模板的智能填充
两种生成策略：批量生成 (1-to-N) 和 聚合生成 (N-to-1)
"""

import re
import io
import zipfile
from typing import Dict, List, Optional, Tuple, BinaryIO
from copy import deepcopy

import pandas as pd
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet

# 尝试导入 python-docx（可选依赖）
try:
    from docx import Document
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from research_models import flatten_profile_for_template, get_research_profile_by_id


# ==================== 占位符处理 ====================

# 支持的占位符格式: {{name}}, {{姓名}}, {{TABLE:xxx}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
TABLE_MARKER_PATTERN = re.compile(r'\{\{TABLE:(\w+)\}\}')


def find_placeholders(text: str) -> List[str]:
    """从文本中提取所有占位符"""
    if not text:
        return []
    return PLACEHOLDER_PATTERN.findall(text)


def replace_placeholders(text: str, data: Dict) -> str:
    """
    替换文本中的占位符
    
    Args:
        text: 包含 {{placeholder}} 的文本
        data: 替换数据字典
    
    Returns:
        替换后的文本
    """
    if not text:
        return text
    
    def replacer(match):
        key = match.group(1).strip()
        # 支持嵌套访问，如 {{contact.phone}}
        if '.' in key:
            parts = key.split('.')
            value = data
            for part in parts:
                if isinstance(value, dict):
                    value = value.get(part, '')
                else:
                    value = ''
                    break
        else:
            value = data.get(key, '')
        
        # 转换为字符串
        if value is None:
            return ''
        return str(value)
    
    return PLACEHOLDER_PATTERN.sub(replacer, text)


# ==================== Excel 处理 ====================

def process_excel_template(
    template_file: BinaryIO,
    profiles: List[Dict],
    strategy: str = "batch"
) -> Tuple[bytes, str]:
    """
    处理 Excel 模板
    
    Args:
        template_file: Excel 模板文件
        profiles: 档案列表
        strategy: "batch" (批量生成) 或 "aggregate" (聚合生成)
    
    Returns:
        (生成的文件内容, 文件名)
    """
    if strategy == "batch":
        return _excel_batch_generate(template_file, profiles)
    else:
        return _excel_aggregate_generate(template_file, profiles)


def _excel_batch_generate(template_file: BinaryIO, profiles: List[Dict]) -> Tuple[bytes, str]:
    """
    批量生成策略：每人一个文件，打包为 ZIP
    """
    output = io.BytesIO()
    
    with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
        for profile in profiles:
            # 重置文件指针
            template_file.seek(0)
            
            # 加载工作簿
            wb = load_workbook(template_file)
            
            # 获取扁平化数据
            flat_data = flatten_profile_for_template(profile)
            
            # 处理每个工作表
            for sheet in wb.worksheets:
                _fill_excel_sheet(sheet, flat_data)
            
            # 保存到内存
            file_buffer = io.BytesIO()
            wb.save(file_buffer)
            file_buffer.seek(0)
            
            # 添加到 ZIP
            name = profile.get("姓名", profile.get("id", "unknown"))
            zf.writestr(f"{name}_filled.xlsx", file_buffer.getvalue())
    
    output.seek(0)
    return output.getvalue(), "batch_generated.zip"


def _excel_aggregate_generate(template_file: BinaryIO, profiles: List[Dict]) -> Tuple[bytes, str]:
    """
    聚合生成策略：所有人填入同一文件的表格中
    """
    template_file.seek(0)
    wb = load_workbook(template_file)
    
    for sheet in wb.worksheets:
        # 查找表格标记行
        table_row = _find_table_marker_row(sheet)
        
        if table_row:
            # 找到表格标记，填充多行数据
            _fill_excel_table(sheet, table_row, profiles)
        else:
            # 没有表格标记，只填充第一个人的数据（如果有的话）
            if profiles:
                flat_data = flatten_profile_for_template(profiles[0])
                _fill_excel_sheet(sheet, flat_data)
    
    # 保存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return output.getvalue(), "aggregated_filled.xlsx"


def _fill_excel_sheet(sheet: Worksheet, data: Dict):
    """填充 Excel 工作表中的占位符"""
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                if '{{' in cell.value:
                    cell.value = replace_placeholders(cell.value, data)


def _find_table_marker_row(sheet: Worksheet) -> Optional[int]:
    """查找包含 {{TABLE:xxx}} 标记的行"""
    for row_idx, row in enumerate(sheet.iter_rows(), start=1):
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                if TABLE_MARKER_PATTERN.search(cell.value):
                    return row_idx
    return None


def _fill_excel_table(sheet: Worksheet, marker_row: int, profiles: List[Dict]):
    """
    在 Excel 表格中填充多人数据
    
    marker_row: 表格标记所在行（将作为模板行）
    """
    # 获取模板行的格式和占位符
    template_cells = []
    for cell in sheet[marker_row]:
        template_cells.append({
            "value": cell.value,
            "column": cell.column
        })
    
    # 清除标记行的 TABLE 标记
    for cell in sheet[marker_row]:
        if cell.value and isinstance(cell.value, str):
            cell.value = TABLE_MARKER_PATTERN.sub('', cell.value).strip()
    
    # 为每个人插入一行数据
    for idx, profile in enumerate(profiles):
        target_row = marker_row + idx
        flat_data = flatten_profile_for_template(profile)
        
        # 如果不是第一行，需要插入新行
        if idx > 0:
            sheet.insert_rows(target_row)
        
        # 填充数据
        for template in template_cells:
            col = template["column"]
            value = template["value"]
            
            if value and isinstance(value, str):
                # 移除 TABLE 标记并替换占位符
                value = TABLE_MARKER_PATTERN.sub('', value)
                value = replace_placeholders(value, flat_data)
            
            sheet.cell(row=target_row, column=col, value=value)


# ==================== Word 处理 ====================

def process_word_template(
    template_file: BinaryIO,
    profiles: List[Dict],
    strategy: str = "batch"
) -> Tuple[bytes, str]:
    """
    处理 Word 模板
    
    Args:
        template_file: Word 模板文件
        profiles: 档案列表
        strategy: "batch" (批量生成) 或 "aggregate" (聚合生成)
    
    Returns:
        (生成的文件内容, 文件名)
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装，无法处理 Word 文件")
    
    if strategy == "batch":
        return _word_batch_generate(template_file, profiles)
    else:
        return _word_aggregate_generate(template_file, profiles)


def _word_batch_generate(template_file: BinaryIO, profiles: List[Dict]) -> Tuple[bytes, str]:
    """批量生成 Word 文件"""
    output = io.BytesIO()
    
    with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
        for profile in profiles:
            template_file.seek(0)
            doc = Document(template_file)
            
            flat_data = flatten_profile_for_template(profile)
            _fill_word_document(doc, flat_data)
            
            file_buffer = io.BytesIO()
            doc.save(file_buffer)
            file_buffer.seek(0)
            
            name = profile.get("姓名", profile.get("id", "unknown"))
            zf.writestr(f"{name}_filled.docx", file_buffer.getvalue())
    
    output.seek(0)
    return output.getvalue(), "batch_generated.zip"


def _word_aggregate_generate(template_file: BinaryIO, profiles: List[Dict]) -> Tuple[bytes, str]:
    """聚合生成 Word 文件"""
    template_file.seek(0)
    doc = Document(template_file)
    
    # 查找并填充表格
    table_filled = False
    for table in doc.tables:
        if _is_template_table(table):
            _fill_word_table(table, profiles)
            table_filled = True
            break
    
    # 如果没有表格，填充第一个人的数据
    if not table_filled and profiles:
        flat_data = flatten_profile_for_template(profiles[0])
        _fill_word_document(doc, flat_data)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue(), "aggregated_filled.docx"


def _fill_word_document(doc, data: Dict):
    """填充 Word 文档中的占位符"""
    # 处理段落
    for paragraph in doc.paragraphs:
        if '{{' in paragraph.text:
            for run in paragraph.runs:
                if '{{' in run.text:
                    run.text = replace_placeholders(run.text, data)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{' in paragraph.text:
                        for run in paragraph.runs:
                            if '{{' in run.text:
                                run.text = replace_placeholders(run.text, data)


def _is_template_table(table) -> bool:
    """判断表格是否包含模板占位符"""
    for row in table.rows:
        for cell in row.cells:
            if '{{' in cell.text:
                return True
    return False


def _fill_word_table(table, profiles: List[Dict]):
    """
    填充 Word 表格（聚合模式）
    假设第一行是表头，第二行是模板行
    """
    if len(table.rows) < 2:
        return
    
    # 获取模板行
    template_row = table.rows[1]
    template_cells = [cell.text for cell in template_row.cells]
    
    # 填充第一个人到模板行
    if profiles:
        flat_data = flatten_profile_for_template(profiles[0])
        for idx, cell in enumerate(template_row.cells):
            cell.text = replace_placeholders(template_cells[idx], flat_data)
    
    # 为其余人添加新行
    for profile in profiles[1:]:
        flat_data = flatten_profile_for_template(profile)
        new_row = table.add_row()
        for idx, cell in enumerate(new_row.cells):
            cell.text = replace_placeholders(template_cells[idx], flat_data)


# ==================== 主入口函数 ====================

def generate_filled_forms(
    template_file: BinaryIO,
    template_filename: str,
    profile_ids: List[str],
    strategy: str = "batch"
) -> Tuple[bytes, str, List[str]]:
    """
    生成填充后的表单
    
    Args:
        template_file: 模板文件对象
        template_filename: 模板文件名（用于判断类型）
        profile_ids: 要填充的档案 ID 列表
        strategy: "batch" 或 "aggregate"
    
    Returns:
        (文件内容, 文件名, 错误列表)
    """
    errors = []
    
    # 获取档案数据
    profiles = []
    for pid in profile_ids:
        profile = get_research_profile_by_id(pid)
        if profile:
            profiles.append(profile)
        else:
            errors.append(f"未找到档案 ID: {pid}")
    
    if not profiles:
        return None, None, ["没有有效的档案数据"]
    
    # 根据文件类型处理
    filename_lower = template_filename.lower()
    
    try:
        if filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls'):
            content, filename = process_excel_template(template_file, profiles, strategy)
        elif filename_lower.endswith('.docx'):
            content, filename = process_word_template(template_file, profiles, strategy)
        else:
            return None, None, [f"不支持的文件格式: {template_filename}"]
        
        return content, filename, errors
    
    except Exception as e:
        return None, None, [f"生成失败: {str(e)}"]


def get_template_placeholders(template_file: BinaryIO, template_filename: str) -> List[str]:
    """
    分析模板文件，提取所有占位符
    
    用于预览模板需要哪些字段
    """
    placeholders = set()
    filename_lower = template_filename.lower()
    
    try:
        template_file.seek(0)
        
        if filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls'):
            wb = load_workbook(template_file)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            found = find_placeholders(cell.value)
                            placeholders.update(found)
        
        elif filename_lower.endswith('.docx'):
            if not DOCX_AVAILABLE:
                return ["(需要安装 python-docx)"]
            
            doc = Document(template_file)
            
            # 段落
            for paragraph in doc.paragraphs:
                found = find_placeholders(paragraph.text)
                placeholders.update(found)
            
            # 表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        found = find_placeholders(cell.text)
                        placeholders.update(found)
    
    except Exception as e:
        return [f"分析失败: {str(e)}"]
    
    return sorted(list(placeholders))


# ==================== 预设模板 ====================

def get_available_field_mappings() -> Dict[str, str]:
    """
    获取所有可用的字段映射
    
    返回 {占位符: 说明}
    """
    return {
        # 基本信息
        "{{name}}": "姓名",
        "{{姓名}}": "姓名",
        "{{phone}}": "电话",
        "{{电话}}": "电话",
        "{{email}}": "邮箱",
        "{{邮箱}}": "邮箱",
        
        # 教育信息
        "{{degree}}": "最高学位",
        "{{学位}}": "最高学位",
        "{{institution}}": "毕业院校",
        "{{院校}}": "毕业院校",
        "{{major}}": "专业",
        "{{专业}}": "专业",
        
        # 论文统计
        "{{publications_count}}": "论文总数",
        "{{论文总数}}": "论文总数",
        "{{sci_count}}": "SCI 论文数",
        "{{SCI论文数}}": "SCI 论文数",
        "{{ei_count}}": "EI 论文数",
        
        # 项目统计
        "{{grants_count}}": "项目总数",
        "{{项目总数}}": "项目总数",
        "{{grants_as_pi}}": "主持项目数",
        "{{主持项目数}}": "主持项目数",
        "{{total_grant_budget}}": "累计项目经费",
        
        # 特殊标记
        "{{TABLE:xxx}}": "表格区域标记（聚合模式）"
    }
